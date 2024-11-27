# File: src/main.py
import streamlit as st
import os
import tempfile
from src.file_handler import FileHandler
from src.redaction_engine import RedactionEngine
from src.ai_suggestions import AIRedactionSuggester
from src.report_generator import RedactionReportGenerator
from src.security_utils import SecurityManager
import docx
import io
from fpdf import FPDF



class MultiDocRedactionApp:
    def __init__(self):
        # Initialize session state if it doesn't exist
        if 'custom_redactions' not in st.session_state:
            st.session_state.custom_redactions = []
        if 'contextual_matches' not in st.session_state:
            st.session_state.contextual_matches = []
        if 'document_text' not in st.session_state:
            st.session_state.document_text = ""
        if 'text_to_redact' not in st.session_state:
            st.session_state.text_to_redact = ""
        if 'reason' not in st.session_state:
            st.session_state.reason = ""
        if 'ai_suggestions' not in st.session_state:
            st.session_state.ai_suggestions = []
        if 'has_generated_suggestions' not in st.session_state:
            st.session_state.has_generated_suggestions = False
        
        self.ai_suggester = AIRedactionSuggester()
        self.redaction_engine = RedactionEngine()
        self.report_generator = RedactionReportGenerator()
        self.security_manager = SecurityManager()
        self.file_handler = FileHandler()

    def show_sample_files(self):
        st.subheader("Sample Test Files")
        st.markdown("""
        Download these sample files to test the redaction capabilities:
        
        Each file contains the same content with various types of sensitive information:
        - Personal Information (names, addresses, emails, phone numbers, SSN)
        - Financial Information (credit card numbers)
        - Work-related Information
        """)

        sample_text = """John Doe lives at 123 Elm Street, Springfield.
His email address is john.doe@example.com, and his phone number is +1-555-123-4567.
Credit Card Number: 4111 1111 1111 1111. Social Security Number: 123-45-6789.

For work-related matters, contact him at work.email@example.com or at (555) 987-6543.
He frequently shops online and uses PayPal with the email paypal.john.doe@example.com."""

        # Create download buttons for each format
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.download_button(
                label="📄 Download TXT Sample",
                data=sample_text,
                file_name="sample_document.txt",
                mime="text/plain"
            )
        
        with col2:
            # Create a simple DOCX in memory
            doc = docx.Document()
            doc.add_heading('Sample Document', 0)
            for paragraph in sample_text.split('\n'):
                doc.add_paragraph(paragraph)
            
            # Save to bytes
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            
            st.download_button(
                label="📘 Download DOCX Sample",
                data=docx_bytes,
                file_name="sample_document.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        with col3:
            # Create PDF in memory
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for line in sample_text.split('\n'):
                pdf.cell(0, 10, txt=line, ln=True)
            
            # Use io.BytesIO to create a proper byte stream
            pdf_bytes_io = io.BytesIO()
            pdf.output(pdf_bytes_io)
            pdf_bytes_io.seek(0)
            
            st.download_button(
                label="📕 Download PDF Sample",
                data=pdf_bytes_io,
                file_name="sample_document.pdf",
                mime="application/pdf"
            )

    def run(self):
        st.title("MultiDoc Redaction Assistant")
        
        self.show_sample_files()
        
        uploaded_file = st.file_uploader(
            "Upload a document",
            type=["pdf", "docx", "txt"],
            help="Supported formats: PDF, DOCX, TXT"
        )
        
        if uploaded_file:
            # Create temporary file
            temp_file_path = os.path.join(tempfile.gettempdir(), uploaded_file.name)
            with open(temp_file_path, "wb") as temp_file:
                temp_file.write(uploaded_file.getvalue())
            
            # Extract and display text
            try:
                document_text = self.file_handler.extract_text(temp_file_path)
                st.subheader("Document Preview")
                st.text(document_text)
                
                # Add AI suggestion button
                if st.button("Get AI Suggestions"):
                    with st.spinner("Analyzing document for sensitive information..."):
                        st.session_state.ai_suggestions = self.ai_suggester.get_redaction_suggestions(
                            temp_file_path,
                            sensitivity=50
                        )
                        st.session_state.has_generated_suggestions = True
                        st.rerun()
                
                # Display AI suggestions (move outside the button click condition)
                if st.session_state.has_generated_suggestions and st.session_state.ai_suggestions:
                    st.subheader("AI-Suggested Redactions")
                    for i, suggestion in enumerate(st.session_state.ai_suggestions):
                        with st.expander(f"Suggestion #{i+1}: {suggestion['type']}"):
                            st.markdown(f"""
                            - **Text to Redact:** `{suggestion['text']}`
                            - **Type:** {suggestion['type']}
                            - **Confidence:** {suggestion['confidence']}%
                            - **Reason:** {suggestion['reason']}
                            """)
                            if st.button("Accept Suggestion", key=f"accept_{i}"):
                                if suggestion['text'] not in [r['text'] for r in st.session_state.custom_redactions]:
                                    st.session_state.custom_redactions.append(suggestion)
                                    st.success("Suggestion added to redaction list!")
                                else:
                                    st.warning("This text is already in the redaction list!")
                elif st.session_state.has_generated_suggestions:
                    st.info("No sensitive information detected by AI.")
                
                # Redaction interface
                st.subheader("Custom Redaction")
                text_to_redact = st.text_input("Enter text to redact:")
                redaction_type = st.selectbox(
                    "Redaction Type:",
                    ["PII", "CREDENTIALS", "FINANCIAL", "CUSTOM"]
                )
                reason = st.text_input("Reason for redaction (optional):")
                
                submit_button = st.button("Add Redaction")
                
                # Debug print
                st.write(f"Current redactions count before processing: {len(st.session_state.custom_redactions)}")
                
                if submit_button and text_to_redact:
                    with st.spinner("Analyzing document for contextual matches..."):
                        # Store matches in session state
                        st.session_state.contextual_matches = self.ai_suggester.analyze_contextual_meaning(
                            document_text,
                            text_to_redact,
                            redaction_type
                        )
                        st.session_state.current_text = text_to_redact
                        st.session_state.current_type = redaction_type
                        st.session_state.current_reason = reason

                # Display contextual matches (move this outside the if block)
                if hasattr(st.session_state, 'contextual_matches') and st.session_state.contextual_matches:
                    st.subheader("Contextual Matches Found:")
                    for idx, match in enumerate(st.session_state.contextual_matches):
                        with st.expander(f"Match #{idx + 1}: {match['text']}"):
                            st.markdown(f"""
                            - **Found Text:** `{match['text']}`
                            - **Confidence:** {match['confidence']}%
                            - **Reason:** {match['reason']}
                            """)
                            
                            # Create unique key for each button
                            button_key = f"add_match_{idx}_{hash(match['text'])}"
                            
                            if st.button("Add this match", key=button_key):
                                new_redaction = {
                                    'text': match['text'],
                                    'type': st.session_state.current_type,
                                    'confidence': match['confidence'],
                                    'reason': st.session_state.current_reason if st.session_state.current_reason else match['reason']
                                }
                                
                                if 'custom_redactions' not in st.session_state:
                                    st.session_state.custom_redactions = []
                                
                                # Check if this exact redaction is already in the list
                                if new_redaction not in st.session_state.custom_redactions:
                                    st.session_state.custom_redactions.append(new_redaction)
                                    st.success(f"Added redaction for: {match['text']}")

                # Display current redactions
                if st.session_state.custom_redactions:
                    st.subheader("Current Redactions")
                    for idx, redaction in enumerate(st.session_state.custom_redactions):
                        with st.expander(f"Redaction #{idx + 1}: {redaction['text']}"):
                            st.markdown(f"""
                            - **Text:** `{redaction['text']}`
                            - **Type:** {redaction['type']}
                            - **Confidence:** {redaction['confidence']}%
                            - **Reason:** {redaction['reason']}
                            """)
                            
                            if st.button("Remove", key=f"remove_{idx}"):
                                st.session_state.custom_redactions.pop(idx)
                                st.rerun()
                
                # Process redactions
                if st.button("Apply Redactions"):
                    # Debug print
                    st.write(f"Attempting to process {len(st.session_state.custom_redactions)} redactions")
                    
                    if not st.session_state.custom_redactions:
                        st.error("Please add at least one redaction before processing")
                    else:
                        try:
                            with st.spinner("Applying redactions..."):
                                output_path = os.path.join(
                                    tempfile.gettempdir(),
                                    f"redacted_{uploaded_file.name}"
                                )
                                
                                redacted_path = self.redaction_engine.redact_document(
                                    temp_file_path,
                                    st.session_state.custom_redactions,  # Use session state
                                    output_path
                                )
                                
                                report = self.report_generator.create_report(
                                    temp_file_path,
                                    redacted_path,
                                    st.session_state.custom_redactions
                                )
                                
                                st.success("Document redacted successfully!")
                                
                                # Create two columns for side-by-side preview
                                col1, col2 = st.columns(2)

                                with col1:
                                    st.subheader("Original Document")
                                    st.text(document_text)

                                with col2:
                                    st.subheader("Redacted Document")
                                    # Extract and display redacted text
                                    redacted_text = self.file_handler.extract_text(redacted_path)
                                    st.text(redacted_text)

                                # Download buttons
                                with open(redacted_path, "rb") as f:
                                    st.download_button(
                                        "Download Redacted Document",
                                        f,
                                        file_name=f"redacted_{uploaded_file.name}"
                                    )
                                
                                st.download_button(
                                    "Download Redaction Report",
                                    report,
                                    file_name="redaction_report.txt"
                                )
                                
                        except Exception as e:
                            st.error(f"Error during redaction: {str(e)}")
                            
                        finally:
                            if os.path.exists(temp_file_path):
                                self.security_manager.secure_file_deletion(temp_file_path)
                            if os.path.exists(output_path):
                                self.security_manager.secure_file_deletion(output_path)
                
            except Exception as e:
                st.error(f"Error processing document: {str(e)}")

def run_app():
    app = MultiDocRedactionApp()
    app.run()

if __name__ == "__main__":
    run_app()
