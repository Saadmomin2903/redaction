# File: src/redaction_engine.py
# Standard library imports
import os
import re
from typing import List, Dict, Tuple, Any
from statistics import median, mean
from collections import Counter

# Third-party library imports
import numpy as np
from sklearn.cluster import KMeans
from sklearn.metrics import silhouette_score
import pdfplumber
import pikepdf
import docx
from docx.oxml import OxmlElement
import fitz


# Local application imports
from .ai_suggestions import AIRedactionSuggester



class AdvancedPDFTextExtractor:
    def __init__(self, pdf_path: str):
        """
        Initialize PDF text extractor with advanced layout analysis
        
        Args:
            pdf_path (str): Path to PDF file
        """
        self.pdf_path = pdf_path
        self.page_layouts = []
    
    def analyze_document_layout(self) -> Dict[str, Any]:
        """
        Perform comprehensive document layout analysis
        
        Returns:
            Dict containing detailed document structure insights
        """
        document_analysis = {
            'total_pages': 0,
            'text_columns': [],
            'font_distribution': {},
            'text_blocks': [],
            'layout_complexity': 0
        }
        
        with pdfplumber.open(self.pdf_path) as pdf:
            document_analysis['total_pages'] = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages):
                # Advanced word extraction with enhanced parameters
                words = page.extract_words(
                    x_tolerance=2,
                    y_tolerance=2,
                    keep_blank_chars=False,
                    use_text_flow=True,
                    split_at_punctuation=True,
                    extra_attrs=['fontname', 'size', 'color']
                )
                
                # Analyze text layout
                page_layout = self._analyze_page_layout(words)
                self.page_layouts.append(page_layout)
                
                # Update document analysis
                document_analysis['text_columns'].extend(page_layout['columns'])
                document_analysis['font_distribution'].update(
                    self._analyze_font_distribution(words)
                )
                document_analysis['text_blocks'].extend(
                    self._extract_text_blocks(words)
                )
                
                # Estimate layout complexity
                document_analysis['layout_complexity'] += self._calculate_layout_complexity(page_layout)
        
        return document_analysis
    
    def _analyze_page_layout(self, words: List[Dict]) -> Dict[str, Any]:
        """
        Analyze page layout with advanced detection techniques
        
        Args:
            words (List[Dict]): Extracted words with detailed attributes
        
        Returns:
            Dict with page layout characteristics
        """
        # Cluster text into potential columns
        columns = self._detect_columns(words)
        
        # Analyze text direction and flow
        text_direction = self._detect_text_direction(words)
        
        # Estimate reading zones
        reading_zones = self._identify_reading_zones(words)
        
        return {
            'columns': columns,
            'text_direction': text_direction,
            'reading_zones': reading_zones,
            'word_count': len(words),
            'vertical_spread': self._calculate_vertical_spread(words)
        }
    
    def _detect_columns(self, words: List[Dict], max_columns: int = 3) -> List[Dict]:
        """
        Detect text columns using advanced clustering
        
        Args:
            words (List[Dict]): Extracted words
            max_columns (int): Maximum number of columns to detect
        
        Returns:
            List of detected column regions
        """
        # Use X-coordinate clustering
        x_coords = [word['x0'] for word in words]
        
        
        
        try:
            # Reshape coordinates for clustering
            X = np.array(x_coords).reshape(-1, 1)
            
            
            best_n_clusters = 1
            best_score = -1
            
            for n in range(1, min(max_columns + 1, len(X))):
                kmeans = KMeans(n_clusters=n, random_state=42)
                cluster_labels = kmeans.fit_predict(X)
                
                # Avoid division by zero
                if len(np.unique(cluster_labels)) > 1:
                    score = silhouette_score(X, cluster_labels)
                    if score > best_score:
                        best_score = score
                        best_n_clusters = n
            
            # Perform clustering with best number of clusters
            kmeans = KMeans(n_clusters=best_n_clusters, random_state=42)
            cluster_labels = kmeans.fit_predict(X)
            
            # Group columns
            columns = []
            for cluster in range(best_n_clusters):
                cluster_words = [word for word, label in zip(words, cluster_labels) if label == cluster]
                columns.append({
                    'x_range': (
                        min(word['x0'] for word in cluster_words),
                        max(word['x1'] for word in cluster_words)
                    ),
                    'words': cluster_words
                })
            
            return columns
        
        except Exception as e:
            print(f"Column detection error: {e}")
            return []
    
    def _detect_text_direction(self, words: List[Dict]) -> str:
        """
        Detect predominant text direction
        
        Args:
            words (List[Dict]): Extracted words
        
        Returns:
            str: Detected text direction ('ltr', 'rtl', 'mixed')
        """
        # Analyze word progression
        word_progressions = [word['x1'] > word['x0'] for word in words]
        
        ltr_ratio = sum(word_progressions) / len(word_progressions) if word_progressions else 0
        
        if ltr_ratio > 0.9:
            return 'ltr'
        elif ltr_ratio < 0.1:
            return 'rtl'
        else:
            return 'mixed'
    
    def _identify_reading_zones(self, words: List[Dict]) -> List[Dict]:
        """
        Identify logical reading zones in the document
        
        Args:
            words (List[Dict]): Extracted words
        
        Returns:
            List of reading zone dictionaries
        """
        # Group words into vertical zones
        sorted_words = sorted(words, key=lambda w: w['top'])
        
        zones = []
        current_zone = None
        
        for word in sorted_words:
            # Start a new zone or extend existing zone
            if not current_zone or word['top'] - current_zone['bottom'] > 10:
                if current_zone:
                    zones.append(current_zone)
                current_zone = {
                    'top': word['top'],
                    'bottom': word['bottom'],
                    'words': [word],
                    'text': word['text']
                }
            else:
                # Extend current zone
                current_zone['bottom'] = max(current_zone['bottom'], word['bottom'])
                current_zone['words'].append(word)
                current_zone['text'] += ' ' + word['text']
        
        if current_zone:
            zones.append(current_zone)
        
        return zones
    
    def _analyze_font_distribution(self, words: List[Dict]) -> Dict[str, int]:
        """
        Analyze font characteristics distribution
        
        Args:
            words (List[Dict]): Extracted words
        
        Returns:
            Dict of font distribution metrics
        """
        font_sizes = [word.get('size', 0) for word in words]
        font_names = [word.get('fontname', 'unknown') for word in words]
        
        return {
            'size_stats': {
                'min': min(font_sizes) if font_sizes else 0,
                'max': max(font_sizes) if font_sizes else 0,
                'median': median(font_sizes) if font_sizes else 0,
                'mean': mean(font_sizes) if font_sizes else 0
            },
            'font_counts': dict(Counter(font_names))
        }
    
    def _extract_text_blocks(self, words: List[Dict]) -> List[Dict]:
        """
        Extract contiguous text blocks with semantic context
        
        Args:
            words (List[Dict]): Extracted words
        
        Returns:
            List of text block dictionaries
        """
        sorted_words = sorted(words, key=lambda w: (w['top'], w['x0']))
        
        text_blocks = []
        current_block = None
        
        for word in sorted_words:
            # Start a new block or continue existing block
            if (not current_block or 
                abs(word['top'] - current_block['bottom']) > 5 or 
                word['x0'] - current_block['right'] > 20):
                
                # Finalize previous block
                if current_block:
                    text_blocks.append(current_block)
                
                # Start new block
                current_block = {
                    'top': word['top'],
                    'bottom': word['bottom'],
                    'left': word['x0'],
                    'right': word['x1'],
                    'text': word['text'],
                    'words': [word],
                    'font_sizes': [word.get('size', 0)],
                    'fonts': [word.get('fontname', 'unknown')]
                }
            else:
                # Extend current block
                current_block['bottom'] = max(current_block['bottom'], word['bottom'])
                current_block['right'] = max(current_block['right'], word['x1'])
                current_block['text'] += ' ' + word['text']
                current_block['words'].append(word)
                current_block['font_sizes'].append(word.get('size', 0))
                current_block['fonts'].append(word.get('fontname', 'unknown'))
        
        # Add final block
        if current_block:
            text_blocks.append(current_block)
        
        return text_blocks
    
    def _calculate_vertical_spread(self, words: List[Dict]) -> float:
        """
        Calculate vertical text spread
        
        Args:
            words (List[Dict]): Extracted words
        
        Returns:
            float: Vertical spread metric
        """
        if not words:
            return 0
        
        top_words = [word['top'] for word in words]
        bottom_words = [word['bottom'] for word in words]
        
        return max(bottom_words) - min(top_words)
    
    def _calculate_layout_complexity(self, page_layout: Dict) -> float:
        """
        Estimate page layout complexity
        
        Args:
            page_layout (Dict): Page layout analysis
        
        Returns:
            float: Layout complexity score
        """
        complexity = 0
        complexity += len(page_layout.get('columns', [])) * 0.5
        complexity += abs(page_layout.get('vertical_spread', 0)) / 100
        complexity += 1 if page_layout.get('text_direction') == 'mixed' else 0
        
        return complexity
    def create_redaction_annotation(self, word: Dict, height: float, padding: float) -> Dict:
        """
        Create a redaction annotation for a word
        
        Args:
            word (Dict): Word information
            height (float): Page height
            padding (float): Padding around redaction
        
        Returns:
            Dict: Redaction annotation
        """
        return {
            'Type': '/Annot',
            'Subtype': '/Redact',
            'QuadPoints': [
                word['x0'] - padding, height - word['y0'] - padding,
                word['x1'] + padding, height - word['y0'] - padding,
                word['x0'] - padding, height - word['y1'] + padding,
                word['x1'] + padding, height - word['y1'] + padding
            ],
            'Rect': [
                word['x0'] - padding, height - word['y1'] + padding,
                word['x1'] + padding, height - word['y0'] - padding
            ],
            'IC': [0, 0, 0],  # Black fill
            'Color': [0, 0, 0],  # Black border
            'F': 4  # Make annotation visible
        }



class RedactionEngine:
    def __init__(self):
        self.ai_suggester = AIRedactionSuggester()
        self.supported_formats = {
            'pdf': self._redact_pdf,
            'docx': self._redact_docx,
            'txt': self._redact_text
        }

    def redact_document(self, file_path: str, suggestions: List[Dict], output_path: str) -> str:
        """
        Redact document based on user-specified redactions
        
        Args:
            file_path (str): Path to original document
            suggestions (List[Dict]): List of user-specified redactions
            output_path (str): Path for saving redacted document
        
        Returns:
            str: Path to redacted document
        """
        try:
            # Normalize file extension
            file_extension = os.path.splitext(file_path)[1].lower().replace('.', '')
            
            # Validate suggestions
            if not suggestions:
                raise ValueError("No redaction suggestions provided")
            
            # Choose appropriate redaction method
            redaction_method = self.supported_formats.get(file_extension)
            if not redaction_method:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Perform redaction
            return redaction_method(file_path, suggestions, output_path)
                
        except Exception as e:
            raise RuntimeError(f"Document redaction failed: {e}")

    def _redact_pdf(self, file_path: str, suggestions: List[Dict], output_path: str) -> str:
        """
        Redact PDF files with enhanced security
        """
        try:
            # Open PDF document
            doc = fitz.open(file_path)
            
            # Process each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Process each redaction suggestion
                for suggestion in suggestions:
                    text_to_redact = suggestion['text']
                    
                    # Method 1: Exact text search
                    instances = page.search_for(text_to_redact, quads=True)
                    
                    # Method 2: Regular expression search for partial matches
                    text_instances = page.get_text("words")
                    for inst in text_instances:
                        if text_to_redact in inst[4]:  # inst[4] contains the text
                            # Create rectangle from word coordinates
                            rect = fitz.Rect(inst[:4])
                            
                            # Add redaction annotation
                            annot = page.add_redact_annot(rect)
                            annot.set_colors(stroke=(0, 0, 0), fill=(0, 0, 0))
                            annot.update()
                    
                    # Apply exact match redactions
                    for quad in instances:
                        rect = quad.rect  # Convert quad to rectangle
                        # Add some padding
                        padding = 2
                        rect.x0 -= padding
                        rect.y0 -= padding
                        rect.x1 += padding
                        rect.y1 += padding
                        
                        # Create and apply redaction
                        annot = page.add_redact_annot(rect)
                        annot.set_colors(stroke=(0, 0, 0), fill=(0, 0, 0))
                        annot.update()
                
                # Apply all redactions to the page
                page.apply_redactions()
            
            # Save redacted document with encryption
            doc.save(
                output_path,
                encryption=fitz.PDF_ENCRYPT_AES_256,
                owner_pw="owner_password",
                user_pw=None,
                permissions=fitz.PDF_PERM_ACCESSIBILITY | 
                           fitz.PDF_PERM_PRINT |
                           fitz.PDF_PERM_COPY
            )
            doc.close()
            
            # Verify redactions
            if not self._verify_redaction(output_path, suggestions):
                raise RuntimeError("Redaction verification failed")
            
            return output_path
            
        except Exception as e:
            raise RuntimeError(f"PDF redaction failed: {e}")

    def _verify_redaction(self, file_path: str, suggestions: List[Dict]) -> bool:
        """
        Verify that redactions were applied correctly
        """
        try:
            doc = fitz.open(file_path)
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Get text in different formats for thorough verification
                text = page.get_text()
                words = page.get_text("words")
                
                # Check each suggestion
                for suggestion in suggestions:
                    text_to_redact = suggestion['text']
                    
                    # Check full text
                    if text_to_redact in text:
                        doc.close()
                        return False
                    
                    # Check individual words
                    for word in words:
                        if text_to_redact in word[4]:  # word[4] contains the text
                            doc.close()
                            return False
            
            doc.close()
            return True
            
        except Exception as e:
            print(f"Verification error: {e}")
            return False

    def _redact_docx(self, file_path: str, suggestions: List[Dict], output_path: str) -> str:
        """Redact Word documents with robust method"""
        try:
            doc = docx.Document(file_path)
            ai_suggester = AIRedactionSuggester()
            
            # Extract full text for AI analysis
            full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Process each suggestion with AI
            for suggestion in suggestions:
                text_to_redact = suggestion['text']
                semantic_matches = ai_suggester.analyze_text_for_redaction(
                    full_text,
                    text_to_redact,
                    suggestion['type']
                )
                
                # Apply redactions to document
                for paragraph in doc.paragraphs:
                    for match in semantic_matches:
                        matched_text = match['text']
                        if matched_text in paragraph.text:
                            # Split runs at match boundaries
                            new_runs = []
                            for run in paragraph.runs:
                                if matched_text in run.text:
                                    parts = run.text.split(matched_text)
                                    for i, part in enumerate(parts):
                                        if part:
                                            new_run = paragraph.add_run(part)
                                            new_run.font.name = run.font.name
                                            new_run.font.size = run.font.size
                                            new_runs.append(new_run)
                                        if i < len(parts) - 1:
                                            # Add redaction
                                            redact_run = paragraph.add_run('█' * len(matched_text))
                                            redact_run.font.name = run.font.name
                                            redact_run.font.size = run.font.size
                                            new_runs.append(redact_run)
                                else:
                                    new_runs.append(run)
                            
                            # Replace original runs with new ones
                            paragraph._p.remove(*paragraph._p.r_lst)
                            for run in new_runs:
                                paragraph._p.append(run._r)
            
            # Save redacted document
            doc.save(output_path)
            return output_path
            
        except Exception as e:
            raise RuntimeError(f"DOCX redaction failed: {e}")

    def _redact_text(self, file_path: str, suggestions: List[Dict], output_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Sort suggestions by length and position to handle overlapping redactions
            suggestions.sort(key=lambda x: (-len(x['text']), x['text'].lower()))
            
            # Create a list of all positions to redact
            redactions = []
            for suggestion in suggestions:
                text = suggestion['text']
                start = 0
                while True:
                    pos = content.find(text, start)
                    if pos == -1:
                        break
                    redactions.append((pos, pos + len(text), '[REDACTED]'))
                    start = pos + 1
            
            # Sort redactions by start position in reverse order
            redactions.sort(key=lambda x: x[0], reverse=True)
            
            # Apply redactions from end to start to maintain correct positions
            content_chars = list(content)
            for start, end, replacement in redactions:
                content_chars[start:end] = replacement
            
            # Convert back to string
            redacted_content = ''.join(content_chars)
            
            # Write redacted content
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(redacted_content)
            
            return output_path
            
        except Exception as e:
            raise RuntimeError(f"Text redaction failed: {e}")

    
    

